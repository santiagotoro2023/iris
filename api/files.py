"""Image, document and video analysis (SPEC.md Phase 6, pulled forward; §40).

An upload becomes text: a picture is described by the vision model, a document has
its text extracted, a video is sampled into frames and its speech transcribed. That
text travels with the message, so the main model can reason about the file on this
turn and every later one without re-reading it.

Uploads are also kept, so each analyser is a tool the model can call again with a
different question. Looking at a picture is an action, and every action IRiS takes
should announce itself (§35).
"""
import asyncio
import base64
import io
import json
import os
import re
import shutil
import time
from pathlib import Path

import httpx
from fastapi import (APIRouter, Depends, File, Form, HTTPException, Request,
                     UploadFile)
from fastapi.responses import FileResponse, StreamingResponse

import activity
import auth
import settings

OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://ollama:11434")
MAX_BYTES = 20 * 1024 * 1024
MAX_DOC_CHARS = 20_000
# Audio is transcribed in slices this long, so one slow chunk cannot take the whole
# transcription down with it.
TRANSCRIBE_CHUNK_SECONDS = 300

IMAGE_TYPES = {"image/png", "image/jpeg", "image/webp", "image/gif", "image/bmp"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}
VIDEO_SUFFIXES = {".mp4", ".mov", ".mkv", ".avi", ".m4v", ".mpg", ".mpeg"}
AUDIO_SUFFIXES = {".mp3", ".m4a", ".wav", ".ogg", ".oga", ".opus", ".flac", ".aac",
                  ".wma", ".aiff", ".aif"}
# .webm is both. The microphone records webm audio and a browser exports webm video
# under the same extension, so it is the one case the suffix genuinely cannot settle
# and ffprobe is asked instead.
AMBIGUOUS_SUFFIXES = {".webm", ".ogg", ".mkv"}
TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".log", ".ini",
                 ".py", ".js", ".ts", ".sh", ".sql", ".html", ".css", ".xml", ".toml"}

# Uploads live here so an analyser can be run again with a different question.
UPLOADS = Path(os.environ.get("UPLOAD_DIR", "/media/uploads"))
# No fixed cap: the only real limit is somewhere to put it. A margin is kept so an
# upload cannot fill the disk the databases live on.
DISK_MARGIN_BYTES = 5 * 1024 * 1024 * 1024


def _room_for(size: int) -> bool:
    try:
        free = shutil.disk_usage(UPLOADS if UPLOADS.is_dir() else "/").free
    except OSError:
        return True
    return size + DISK_MARGIN_BYTES < free

router = APIRouter(prefix="/files", tags=["files"])

settings.setting(
    "vision.model", type="string",
    enum=["qwen2.5vl:3b", "qwen2.5vl:7b", "llava:7b", "moondream"],
    default=os.environ.get("IRIS_VISION_MODEL", "qwen2.5vl:3b"),
    title="Vision model",
    description="Describes uploaded images. Loading it evicts the language model from "
                "the GPU briefly, so smaller is smoother on 8 GB.")
settings.setting(
    "vision.video_frames", type="integer", minimum=2, maximum=16, default=6,
    title="Frames to look at in a video",
    description="More frames catch more, and each one costs a pass of the vision "
                "model.")
settings.setting(
    "vision.keep_uploads_hours", type="integer", minimum=0, maximum=8760, default=72,
    title="Keep uploaded files for (hours)",
    description="How long a file stays available for follow-up questions. 0 deletes "
                "it as soon as it has been read.")
settings.setting(
    "vision.diarize", type="boolean", default=True,
    title="Label speakers in recordings",
    description="Works out who said what in a video or recording. Runs on the CPU, "
                "so it never takes VRAM from the language model, and adds roughly "
                "half again to the time a transcription takes.")
settings.setting(
    "vision.prompt", type="string", format="multiline",
    default="Describe this image in detail. Include any text you can read verbatim, "
            "and note anything unusual or noteworthy.",
    title="Image prompt",
    description="What IRiS is asked when it looks at a picture.")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _as_png(data: bytes) -> bytes:
    """The vision model rejects webp and friends outright, so normalise first."""
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


async def _describe_image(data: bytes, prompt: str) -> str:
    try:
        data = _as_png(data)
    except Exception:
        pass  # already a format the model accepts, or genuinely unreadable
    body = {"model": settings.get("vision.model"),
            "messages": [{"role": "user",
                          "content": prompt,
                          "images": [base64.b64encode(data).decode()]}],
            "stream": False}
    async with httpx.AsyncClient(timeout=300) as c:
        r = await c.post(f"{OLLAMA_URL}/api/chat", json=body)
    if r.status_code != 200:
        raise HTTPException(502, f"vision model: {r.text[:300]}")
    return (r.json().get("message") or {}).get("content", "").strip()


def kind_of(name: str, ctype: str = "") -> str:
    suffix = ("." + name.lower().rsplit(".", 1)[-1]) if "." in name else ""
    ctype = (ctype or "").split(";")[0]
    if suffix in AMBIGUOUS_SUFFIXES:
        # The browser's own guess first, then the container default. Probing the file
        # would be better and needs the bytes, which `kind_of` does not have; the
        # stored path is probed properly in `_probe_kind` below.
        if ctype.startswith("audio/"):
            return "audio"
        if ctype.startswith("video/"):
            return "video"
        return "video" if suffix == ".mkv" else "audio"
    if suffix in AUDIO_SUFFIXES or ctype.startswith("audio/"):
        return "audio"
    if suffix in VIDEO_SUFFIXES or ctype.startswith("video/"):
        return "video"
    if ctype in IMAGE_TYPES or suffix in IMAGE_SUFFIXES:
        return "image"
    if suffix == ".pdf" or ctype == "application/pdf":
        return "pdf"
    if suffix == ".docx":
        return "document"
    if suffix in TEXT_SUFFIXES or ctype.startswith("text/"):
        return "text"
    return "file"


@router.post("/upload")
async def store(file: UploadFile = File(...),
                user: dict = Depends(auth.active_user)):
    """Keep the file and say what it is. Nothing is read yet.

    Reading a two-minute video takes a minute, and doing it before the message is
    sent means staring at a progress bar with nothing to do. It happens during the
    reply instead, as a tool call with its own banner.
    """
    data = await file.read()
    if not data:
        raise HTTPException(400, "empty file")
    if not _room_for(len(data)):
        raise HTTPException(413, f"no room for {len(data) // 2**20} MB on the disk")
    name = file.filename or "upload"
    path = keep(name, data)
    kind = kind_of(name, file.content_type or "")
    await activity.record("file.upload", f"{name} ({kind}, {len(data) // 1024} KiB)",
                          user["username"])
    return {"name": path.name, "kind": kind, "bytes": len(data)}


MEDIA_TYPES = {".mp4": "video/mp4", ".mov": "video/quicktime",
               ".mkv": "video/x-matroska", ".webm": "video/webm",
               ".avi": "video/x-msvideo", ".m4v": "video/x-m4v",
               ".mp3": "audio/mpeg", ".m4a": "audio/mp4", ".wav": "audio/wav",
               ".ogg": "audio/ogg", ".oga": "audio/ogg", ".opus": "audio/opus",
               ".flac": "audio/flac", ".aac": "audio/aac", ".aiff": "audio/aiff",
               ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
               ".webp": "image/webp", ".gif": "image/gif", ".pdf": "application/pdf"}


@router.get("/media/{name}")
async def media(name: str, request: Request, _: dict = Depends(auth.active_user)):
    """The stored upload itself, so the browser can play it beside its transcript.

    Range requests are answered properly rather than ignored: without them a browser
    will happily play a video from the start and refuse to seek, which makes a
    clickable transcript useless.
    """
    path = find(name)
    if not path:
        raise HTTPException(404, f"no uploaded file matching {name!r}")
    size = path.stat().st_size
    ctype = MEDIA_TYPES.get(path.suffix.lower(), "application/octet-stream")
    common = {"accept-ranges": "bytes", "cache-control": "private, max-age=3600"}

    header = request.headers.get("range", "")
    match = re.match(r"bytes=(\d*)-(\d*)$", header.strip()) if header else None
    if not match:
        return FileResponse(path, media_type=ctype, headers=common)

    first, last = match.group(1), match.group(2)
    if first:
        start = int(first)
        end = int(last) if last else size - 1
    else:
        # "bytes=-500" is the last 500 bytes, not the first 500. Players use it to
        # find the moov atom at the end of an mp4.
        start, end = max(0, size - int(last or 0)), size - 1
    end = min(end, size - 1)
    if start > end or start >= size:
        raise HTTPException(416, "that range is outside the file",
                            headers={"content-range": f"bytes */{size}"})

    def chunks():
        with path.open("rb") as handle:
            handle.seek(start)
            left = end - start + 1
            while left > 0:
                block = handle.read(min(262144, left))
                if not block:
                    break
                left -= len(block)
                yield block

    return StreamingResponse(
        chunks(), status_code=206, media_type=ctype,
        headers={**common, "content-range": f"bytes {start}-{end}/{size}",
                 "content-length": str(end - start + 1)})


@router.get("/transcript/{name}")
async def transcript(name: str, _: dict = Depends(auth.active_user)):
    """The cached extraction, for the clickable timeline in the chat. 404 means the
    file has not been read yet, which is a different thing from having no speech."""
    path = find(name)
    if not path:
        raise HTTPException(404, f"no uploaded file matching {name!r}")
    cached = _read_cache(path)
    if not cached:
        raise HTTPException(404, "that file has not been read yet")
    return {"name": path.name, "kind": cached.get("kind", ""),
            "duration": cached.get("duration", 0),
            "segments": cached.get("segments") or [],
            "scenery": cached.get("scenery") or [],
            "speakers": cached.get("speakers", "none"),
            "failure": cached.get("failure", "")}


@router.post("/analyze")
async def analyze(file: UploadFile = File(...),
                  question: str = Form(""),
                  user: dict = Depends(auth.active_user)):
    data = await file.read()
    if not data:
        raise HTTPException(400, "empty file")
    name = file.filename or "upload"
    lower = name.lower()
    ctype = (file.content_type or "").split(";")[0]
    suffix = "." + lower.rsplit(".", 1)[-1] if "." in lower else ""
    is_video = suffix in VIDEO_SUFFIXES or ctype.startswith("video/")
    if not _room_for(len(data)):
        raise HTTPException(413, f"no room for {len(data) // 2**20} MB on the disk")

    path = keep(name, data)

    try:
        if is_video:
            kind = "video"
            text = await describe_video(path, question)
        elif ctype in IMAGE_TYPES or suffix in IMAGE_SUFFIXES:
            kind = "image"
            # A specific question beats a generic description.
            text = await _describe_image(
                data, question.strip() or settings.get("vision.prompt"))
        elif suffix == ".pdf" or ctype == "application/pdf":
            kind = "pdf"
            text = _extract_pdf(data)
        elif suffix == ".docx":
            kind = "document"
            text = _extract_docx(data)
        elif suffix in TEXT_SUFFIXES or ctype.startswith("text/"):
            kind = "text"
            text = data.decode("utf-8", errors="replace")
        else:
            raise HTTPException(415, f"unsupported file type: {ctype or suffix or '?'}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"could not read {name}: {e}")

    text = (text or "").strip()
    truncated = len(text) > MAX_DOC_CHARS
    if truncated:
        text = text[:MAX_DOC_CHARS] + "\n[...truncated]"
    if not text:
        text = "(no readable content)"

    await activity.record("files.analyze", f"{kind}: {name} ({len(data)//1024} KB)",
                          user["username"])
    return {"name": name, "kind": kind, "bytes": len(data),
            "truncated": truncated, "text": text}


# ---------------------------------------------------------------- video ----

async def _run(*args: str, timeout: int = 300) -> bytes:
    proc = await asyncio.create_subprocess_exec(
        *args, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE)
    try:
        out, err = await asyncio.wait_for(proc.communicate(), timeout)
    except asyncio.TimeoutError:
        if proc.returncode is None:
            try:
                proc.kill()
            except ProcessLookupError:
                pass
        raise HTTPException(504, "that took too long to read")
    if proc.returncode != 0 and not out:
        raise HTTPException(500, err.decode()[:200] or "could not read the file")
    return out


async def _duration(path: Path) -> float:
    try:
        out = await _run("ffprobe", "-v", "error", "-show_entries",
                         "format=duration", "-of", "csv=p=0", str(path), timeout=60)
        return float(out.decode().strip())
    except Exception:
        return 0.0


def _stamp(seconds: float) -> str:
    return f"{int(seconds // 60):02d}:{int(seconds % 60):02d}"


# ------------------------------------------------------------------ cache ----
# Reading a seven minute video costs seven minutes. Doing that again because the
# follow-up question was "and what did he say about the company?" is the difference
# between an assistant and a very slow one. The extraction is kept beside the file.

CACHE_VERSION = 3


def _sidecar(path: Path) -> Path:
    return path.with_suffix(path.suffix + ".iris.json")


def _read_cache(path: Path) -> dict | None:
    """Keyed on size and mtime, so a file replaced under the same name is re-read
    rather than answered from what the old one said."""
    side = _sidecar(path)
    if not side.is_file():
        return None
    try:
        data = json.loads(side.read_text())
        stat = path.stat()
    except (OSError, ValueError):
        return None
    if (data.get("version") != CACHE_VERSION
            or data.get("size") != stat.st_size
            or abs(data.get("mtime", 0) - stat.st_mtime) > 1):
        return None
    return data


def _write_cache(path: Path, data: dict) -> None:
    try:
        stat = path.stat()
        _sidecar(path).write_text(json.dumps(
            {**data, "version": CACHE_VERSION, "size": stat.st_size,
             "mtime": stat.st_mtime}))
    except OSError as e:
        # Worth saying: silently not caching turns into "why is this still slow".
        print(f"[files] could not cache {path.name}: {e}", flush=True)


async def _probe_kind(path: Path) -> str:
    """Audio or video, decided by what streams the file actually has. A .webm from
    the microphone and a .webm from a screen recorder are the same extension and
    very much not the same thing."""
    if path.suffix.lower() not in AMBIGUOUS_SUFFIXES:
        return kind_of(path.name)
    try:
        out = await _run("ffprobe", "-v", "error", "-select_streams", "v:0",
                         "-show_entries", "stream=codec_type", "-of", "csv=p=0",
                         str(path), timeout=30)
        return "video" if out.decode().strip() else "audio"
    except Exception:
        return kind_of(path.name)


# Speaker identity does not survive chunking: whoever the clusterer calls "Speaker 1"
# in the second chunk need not be the same person it called that in the first. Under
# this length the file goes over in one request so the labels mean one thing all the
# way through; over it, the transcript says so rather than implying a continuity that
# is not there.
DIARIZE_WHOLE_SECONDS = 900


async def transcribe_media(path: Path, seconds: float,
                           diarize: bool = False) -> tuple[list[dict], str, str]:
    """The speech, with timestamps and optionally speakers.

    Returns (segments, failure, speaker_scope). Chunked because a whole file is one
    unbounded request: 435 seconds of audio pushed onto the CPU (the GPU being full of
    the language model) ran past the 300-second client timeout, the failure was
    swallowed, and the reply was written from the pictures alone. A chunk is bounded,
    and a chunk that fails costs only its own minutes rather than the whole thing.
    """
    import voice
    total = seconds or TRANSCRIBE_CHUNK_SECONDS
    whole = diarize and total <= DIARIZE_WHOLE_SECONDS
    step = total if whole else TRANSCRIBE_CHUNK_SECONDS
    scope = "whole" if whole else ("per-part" if diarize else "none")

    out: list[dict] = []
    at = 0.0
    while at < total:
        span = min(step, total - at)
        try:
            wav = await _run("ffmpeg", "-nostdin", "-loglevel", "error",
                             "-ss", f"{at:.2f}", "-t", f"{span:.2f}", "-i", str(path),
                             "-vn", "-ac", "1", "-ar", "16000", "-f", "wav", "-",
                             timeout=600)
        except Exception as e:
            return out, f"the audio could not be extracted ({e})", scope
        if not wav or len(wav) < 1000:
            break
        try:
            # Whisper is roughly real-time on the CPU it falls back to, and the
            # diarizer adds a forward pass per segment on top; the margin covers a
            # cold model load as well.
            result = await voice.transcribe_bytes(
                wav, "media.wav", "audio/wav",
                timeout=max(300.0, span * (8 if diarize else 4)),
                timestamps=True, diarize=diarize)
        except Exception as e:
            note = (f"the audio after {_stamp(at)} could not be transcribed ({e})"
                    if out else f"the audio could not be transcribed ({e})")
            return out, note, scope
        if result.get("diarization"):
            print(f"[files] speakers unavailable: {result['diarization']}", flush=True)
            scope = "none"
        part = int(at // step) + 1
        for segment in result.get("segments") or []:
            text = (segment.get("text") or "").strip()
            if not text:
                continue
            who = segment.get("speaker") or ""
            if who and scope == "per-part":
                who = f"{who} (part {part})"
            out.append({"start": round(at + segment.get("start", 0), 2),
                        "end": round(at + segment.get("end", 0), 2),
                        "text": text, "speaker": who})
        if not result.get("segments"):
            body = (result.get("text") or "").strip()
            if body:
                out.append({"start": at, "end": at + span, "text": body,
                            "speaker": ""})
        at += span
    if not any(s.get("speaker") for s in out):
        scope = "none"
    return out, "", scope


def transcript_lines(segments: list[dict]) -> list[str]:
    return [f"[{_stamp(s['start'])}] "
            + (f"{s['speaker']}: " if s.get("speaker") else "")
            + s["text"] for s in segments]


def _render(cached: dict) -> str:
    """The cached extraction as the model reads it. Speech first: it is almost always
    the point of an uploaded recording, and six sentences of scene description ahead
    of it were enough to make the model answer about the background instead."""
    kind = cached.get("kind", "video")
    seconds = cached.get("duration") or 0
    parts = []
    if seconds:
        parts.append(f"{kind.title()}, {int(seconds // 60)}m {int(seconds % 60)}s "
                     f"long.")
    segments = cached.get("segments") or []
    if segments:
        body = "\n".join(transcript_lines(segments))
        if len(body) > MAX_DOC_CHARS:
            body = body[:MAX_DOC_CHARS] + "\n[...transcript truncated]"
        header = f"TRANSCRIPT (what is said, this is the content of the {kind})"
        if cached.get("speakers") == "per-part":
            header += (". This recording was transcribed in parts and the speaker "
                       "numbering restarts in each part, so 'Speaker 1 (part 2)' is "
                       "not necessarily 'Speaker 1 (part 1)'")
        parts.append(f"{header}:\n{body}")
    elif cached.get("failure"):
        parts.append(f"No transcript: {cached['failure']}. Do not guess at what was "
                     f"said.")
    else:
        parts.append(f"No speech in this {kind}.")
    scenery = cached.get("scenery") or []
    if scenery:
        parts.append("SCENERY (what the camera shows, background only):\n"
                     + "\n".join(f"[{_stamp(s['at'])}] {s['text']}" for s in scenery))
    return "\n\n".join(parts) or f"(nothing readable in that {kind})"


async def read_recording(path: Path, question: str = "",
                         refresh: bool = False) -> str:
    """A video or a recording, read once and kept.

    Frames are sampled evenly rather than from the first N seconds: a video's subject
    is rarely in its opening frame, and consecutive frames say almost the same thing.
    """
    if not refresh:
        cached = _read_cache(path)
        if cached:
            return _render(cached)

    kind = await _probe_kind(path)
    seconds = await _duration(path)
    diarize = settings.get("vision.diarize")
    segments, failure, speakers = await transcribe_media(path, seconds, diarize)
    if failure:
        # Said out loud rather than only logged. Silence here reads as "nothing was
        # said", and the model then answers confidently from the scenery.
        print(f"[files] {path.name}: {failure}", flush=True)

    described = []
    if kind == "video":
        frames = settings.get("vision.video_frames")
        # The image prompt asks for detail, which is right for one picture and useless
        # for six: it produced 200 words a frame about colour calibration.
        prompt = ("In one short sentence, say what is visible here. No preamble, no "
                  "explanation, no offer of help.")
        for i in range(frames):
            at = (seconds * (i + 0.5) / frames) if seconds else i * 2
            try:
                jpeg = await _run("ffmpeg", "-nostdin", "-loglevel", "error",
                                  "-ss", f"{at:.2f}", "-i", str(path),
                                  "-frames:v", "1", "-q:v", "4", "-f", "image2", "-",
                                  timeout=120)
                if not jpeg:
                    continue
                text = await _describe_image(jpeg, prompt)
            except Exception as e:
                print(f"[files] frame at {at:.1f}s failed: {e}", flush=True)
                continue
            one_line = " ".join(text.split())[:220]
            if described and one_line == described[-1]["text"]:
                continue          # a static shot repeats; say it once
            described.append({"at": round(at, 2), "text": one_line})

    cached = {"kind": kind, "duration": seconds, "segments": segments,
              "scenery": described, "failure": failure, "speakers": speakers,
              "text": " ".join(s["text"] for s in segments)}
    _write_cache(path, cached)
    return _render(cached)


# The name the tool and the older call sites use. A recording is read the same way
# whether or not it happens to carry pictures.
describe_video = read_recording
describe_audio = read_recording


# -------------------------------------------------------------- storage ----

_SAFE = re.compile(r"[^A-Za-z0-9._-]+")


def stored_path(name: str) -> Path:
    """Never let an uploaded name escape the upload directory."""
    return UPLOADS / (_SAFE.sub("_", Path(name).name) or "upload")


def keep(name: str, data: bytes) -> Path:
    UPLOADS.mkdir(parents=True, exist_ok=True)
    path = stored_path(name)
    path.write_bytes(data)
    prune()
    return path


def prune() -> None:
    hours = settings.get("vision.keep_uploads_hours")
    if not UPLOADS.is_dir():
        return
    cutoff = time.time() - hours * 3600
    for item in UPLOADS.iterdir():
        if item.is_file() and (hours == 0 or item.stat().st_mtime < cutoff):
            item.unlink(missing_ok=True)


def find(name: str) -> Path | None:
    """Match on the stored name, then on a loose contains, because the model will
    quote the filename back approximately."""
    if not UPLOADS.is_dir():
        return None
    exact = stored_path(name)
    if exact.is_file():
        return exact
    wanted = _SAFE.sub("_", name).lower()
    for item in sorted(UPLOADS.iterdir(),
                       key=lambda p: p.stat().st_mtime, reverse=True):
        if wanted and wanted in item.name.lower():
            return item
    return None


def recent(limit: int = 10) -> list[str]:
    if not UPLOADS.is_dir():
        return []
    return [p.name for p in sorted(UPLOADS.iterdir(),
                                   key=lambda p: p.stat().st_mtime,
                                   reverse=True)[:limit] if p.is_file()]


async def analyse_stored(name: str, question: str = "",
                         refresh: bool = False) -> str:
    """Used by the analysis tools: read a kept upload again, with a new question.

    Everything but an image answers from the cached extraction. A follow-up about a
    seven minute video is a question about its transcript, and re-transcribing it to
    answer costs another seven minutes for text that has not changed. An image is the
    exception: a specific question there is one cheap pass of the vision model and
    genuinely gets a better answer than the generic description would.
    """
    path = find(name)
    if not path:
        have = recent()
        return (f"No uploaded file matching {name!r}. "
                + (f"Available: {', '.join(have)}." if have else "Nothing uploaded."))
    suffix = path.suffix.lower()
    if suffix in VIDEO_SUFFIXES or suffix in AUDIO_SUFFIXES \
            or suffix in AMBIGUOUS_SUFFIXES:
        return await read_recording(path, question, refresh)
    data = path.read_bytes()
    if suffix in IMAGE_SUFFIXES:
        return await _describe_image(
            data, question.strip() or settings.get("vision.prompt"))
    if not refresh:
        cached = _read_cache(path)
        if cached:
            return cached["text"] or "(no readable content)"
    if suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix == ".docx":
        text = _extract_docx(data)
    else:
        text = data.decode("utf-8", errors="replace")
    text = text.strip()[:MAX_DOC_CHARS]
    _write_cache(path, {"kind": kind_of(path.name), "text": text})
    return text or "(no readable content)"
